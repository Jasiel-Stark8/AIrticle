"""Dashboard view logic:
- Send topic and keywords from client to GPT-3.5-turbo
- Generate content

Exports:
- Set root directory that will be created on users device
- Formats {DOCX, PDF, MD, TXT}
"""
import os
from flask import Flask, render_template, request, send_from_directory, abort, session
from werkzeug.utils import secure_filename
import markdown
from api.v1.views.gpt import generate_article
from models.save_article import Article
from database import db
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# BLUEPRINT
from flask import Blueprint
dashboard = Blueprint('dashboard', __name__)

UPLOAD_FOLDER = 'exports/'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md'}


# ======================== Generate Article Logic ========================
@dashboard.route('/generate', methods=['POST'], strict_slashes=False)
def generate_content():
    """Get article parameters from client to feed to GPT"""
    topic = request.form['topic']
    keywords = [k.strip() for k in request.form['keywords'].split(',')]
    article_length = request.form['article_length']

    article = generate_article(topic, keywords, article_length)

    return render_template('dashboard.html', article=article)


# ======================== Export Article Logic ========================
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@dashboard.route('/export', methods=['POST'], strict_slashes=False)
def export_content():
    """Export the content to desired format"""
    content = request.form['content']
    topic = request.form['topic']
    export_format = request.form['format']

    if export_format == 'docx':
        filename = export_docx(content, topic)
    elif export_format == 'pdf':
        filename = export_pdf(content, topic)
    elif export_format == 'txt':
        filename = export_txt(content, topic)
    elif export_format == 'markdown':
        filename = export_markdown(content, topic)
    else:
        return "Unsupported Format", 400

    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)


def export_docx(content, topic):
    """Export content as DOCX"""
    doc = Document()
    doc.add_heading(topic, level=1)
    doc.add_paragraph(content)
    filename = secure_filename(f"{topic}.docx")
    doc.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    return filename


def export_pdf(content, topic):
    """Export content as PDF"""
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(f"{topic}.pdf"))
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 100, content)
    c.save()
    return f"{topic}.pdf"


def export_txt(content, topic):
    """Export content as Txt"""
    filename = secure_filename(f"{topic}.txt")
    with open(os.path.join(app.config['UPLOAD_FOLDER'], filename), 'w') as f:
        f.write(content)
    return filename


def export_markdown(content, topic):
    """Export content as Markdown"""
    filename = secure_filename(f"{topic}.md")
    with open(os.path.join(app.config['UPLOAD_FOLDER'], filename), 'w') as f:
        f.write(f"#{topic}\n\n")
        f.write(content)
    return filename


@dashboard.route('/save_article', methods=['POST'], strict_slashes=False)
def save_article():
    """Save article to database"""
    user_id = session.get('user_id')
    if user_id is None:
        abort(401)
    else:
        topic = request.form['topic']
        content = request.form['content']
        export_format = request.form['format']

        new_article = Article(
            title=topic,
            content=content,
            export_format=export_format
        )
        new_article.user_id = session.id

        db.session.add(new_article)
        db.session.commit()

        return "Article Saved", 200
    # else:
    #     return "Problem saving article, click save again", 400


@dashboard.route('/new_article', methods=['POST'], strict_slashes=False)
def create_new_article():
    """Create new article"""
    return render_template('dashboard.html')
